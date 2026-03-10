"""Deterministic output generation tools for documents and files.

These tools are intentionally native and deterministic:
- HTML pages are saved directly to disk.
- Excel files are generated with xlsxwriter.
- Word documents are generated with python-docx.

They complement privileged sandbox execution instead of replacing it.
Use sandbox/code execution for dynamic computation or browser automation.
Use these tools when the user simply needs a clean deliverable file.
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

from langchain_core.tools import tool

from app.core.generated_files import (
    build_generated_file_url,
    get_generated_dir as _shared_generated_dir,
    guess_generated_media_type,
)
from app.engine.tools.registry import ToolAccess, ToolCategory, get_tool_registry
from app.engine.tools.runtime_context import (
    build_runtime_correlation_metadata,
    emit_tool_bus_event,
    get_current_tool_runtime_context,
)

logger = logging.getLogger(__name__)

def _slugify(value: str, default: str) -> str:
    text = (value or "").strip().lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = re.sub(r"-{2,}", "-", text).strip("-")
    return text[:48] or default


def _timestamp() -> str:
    from time import strftime

    return strftime("%Y%m%d_%H%M%S")


def _json_error(message: str) -> str:
    return json.dumps({"error": message}, ensure_ascii=False)


def get_generated_dir() -> Path:
    """Public accessor for the generated artifact workspace."""
    return _shared_generated_dir()


def _guess_generated_media_type(filepath: Path) -> str:
    return guess_generated_media_type(filepath)


def _artifact_id(prefix: str, filename: str) -> str:
    runtime = get_current_tool_runtime_context()
    seed = (
        runtime.tool_call_id
        if runtime
        else None
    ) or (
        runtime.request_id
        if runtime
        else None
    ) or (
        runtime.session_id
        if runtime
        else None
    ) or _timestamp()
    return f"{prefix}-{seed}-{_slugify(filename, prefix)}"


def _build_artifact_metadata(filepath: Path) -> dict[str, Any]:
    runtime = get_current_tool_runtime_context()
    filename = filepath.name
    metadata = {
        "file_path": str(filepath),
        "file_url": build_generated_file_url(filename),
        "filename": filename,
        "content_type": _guess_generated_media_type(filepath),
        "generated_via": "native_output_tool",
    }
    metadata.update(build_runtime_correlation_metadata(runtime))
    return metadata


def _emit_generated_artifact(
    *,
    artifact_type: str,
    title: str,
    filename: str,
    content: str,
    language: str = "",
    metadata: dict[str, Any],
) -> None:
    emit_tool_bus_event(
        {
            "type": "artifact",
            "content": {
                "artifact_type": artifact_type,
                "artifact_id": _artifact_id(artifact_type, filename),
                "title": title,
                "content": content,
                "language": language,
                "metadata": metadata,
            },
        }
    )


def _normalize_table_rows(rows: Any) -> tuple[list[str], list[dict[str, Any]]]:
    if not isinstance(rows, list) or not rows:
        raise ValueError("rows_json must be a non-empty JSON array of objects")

    normalized_rows: list[dict[str, Any]] = []
    columns: list[str] = []
    seen: set[str] = set()

    for row in rows:
        if not isinstance(row, dict):
            raise ValueError("Each row in rows_json must be a JSON object")
        normalized_rows.append(row)
        for key in row.keys():
            key_text = str(key)
            if key_text not in seen:
                columns.append(key_text)
                seen.add(key_text)

    if not columns:
        raise ValueError("rows_json must contain at least one column")

    return columns, normalized_rows


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return str(value)
    return json.dumps(value, ensure_ascii=False)


def _ensure_full_html(title: str, html_content: str) -> str:
    content = (html_content or "").strip()
    if not content:
        raise ValueError("html_content must not be empty")

    has_html_shell = "<html" in content.lower()
    if has_html_shell:
        return content

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{title}</title>
</head>
<body>
{content}
</body>
</html>
"""


def _write_markdown_to_docx(document: Any, markdown_content: str) -> None:
    """Very small markdown -> docx mapper for headings, bullets, and paragraphs."""
    for raw_line in (markdown_content or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            document.add_paragraph(text, style="List Number")
            continue

        document.add_paragraph(stripped)


@tool(
    description=(
        "Generate a real .html file from self-contained HTML content. "
        "Use when the user asks for a landing page, microsite, or exportable HTML file."
    )
)
def tool_generate_html_file(
    html_content: str,
    title: str = "Wiii Landing Page",
) -> str:
    """Save a self-contained HTML file into the workspace."""
    try:
        rendered = _ensure_full_html(title=title, html_content=html_content)
        output_dir = _shared_generated_dir()
        filename = f"{_slugify(title, 'wiii-page')}_{_timestamp()}.html"
        filepath = output_dir / filename
        filepath.write_text(rendered, encoding="utf-8")
        metadata = _build_artifact_metadata(filepath)
        _emit_generated_artifact(
            artifact_type="html",
            title=title,
            filename=filename,
            content=rendered,
            language="html",
            metadata=metadata,
        )
        return json.dumps(
            {
                "file_path": str(filepath),
                "file_url": metadata["file_url"],
                "filename": filename,
                "format": "html",
                "title": title,
            },
            ensure_ascii=False,
        )
    except Exception as exc:
        logger.error("[OUTPUT_TOOLS] HTML generation failed: %s", exc)
        return _json_error(f"HTML generation failed: {exc}")


@tool(
    description=(
        "Generate a real .xlsx spreadsheet from a JSON array of objects. "
        "Use when the user wants an Excel file, spreadsheet, or structured table export."
    )
)
def tool_generate_excel_file(
    rows_json: str,
    title: str = "Wiii Data Export",
    sheet_name: str = "Sheet1",
) -> str:
    """Create a generic Excel spreadsheet for the user."""
    try:
        rows = json.loads(rows_json)
        columns, normalized_rows = _normalize_table_rows(rows)
    except json.JSONDecodeError as exc:
        return _json_error(f"Invalid JSON: {str(exc)[:100]}")
    except Exception as exc:
        return _json_error(str(exc))

    try:
        import xlsxwriter

        output_dir = _shared_generated_dir()
        filename = f"{_slugify(title, 'wiii-data')}_{_timestamp()}.xlsx"
        filepath = output_dir / filename
        workbook = xlsxwriter.Workbook(str(filepath))
        worksheet = workbook.add_worksheet((sheet_name or "Sheet1")[:31])

        header_fmt = workbook.add_format(
            {
                "bold": True,
                "bg_color": "#1F4E79",
                "font_color": "#FFFFFF",
                "border": 1,
                "text_wrap": True,
            }
        )
        cell_fmt = workbook.add_format({"border": 1, "text_wrap": True, "valign": "top"})

        max_widths = [len(col) for col in columns]
        for col_idx, column in enumerate(columns):
            worksheet.write(0, col_idx, column, header_fmt)

        for row_idx, row in enumerate(normalized_rows, start=1):
            for col_idx, column in enumerate(columns):
                value = row.get(column)
                text = _cell_to_text(value)
                worksheet.write(row_idx, col_idx, text, cell_fmt)
                max_widths[col_idx] = min(max(max_widths[col_idx], len(text)), 48)

        for col_idx, width in enumerate(max_widths):
            worksheet.set_column(col_idx, col_idx, width + 2)

        workbook.close()
        metadata = _build_artifact_metadata(filepath)
        metadata.update(
            {
                "sheet_name": (sheet_name or "Sheet1")[:31],
                "row_count": len(normalized_rows),
                "column_count": len(columns),
                "columns": columns,
            }
        )
        _emit_generated_artifact(
            artifact_type="excel",
            title=title,
            filename=filename,
            content=json.dumps(normalized_rows, ensure_ascii=False),
            metadata=metadata,
        )

        return json.dumps(
            {
                "file_path": str(filepath),
                "file_url": metadata["file_url"],
                "filename": filename,
                "format": "xlsx",
                "title": title,
                "sheet_name": metadata["sheet_name"],
                "row_count": metadata["row_count"],
                "column_count": metadata["column_count"],
                "columns": columns,
            },
            ensure_ascii=False,
        )
    except Exception as exc:
        logger.error("[OUTPUT_TOOLS] Excel generation failed: %s", exc)
        return _json_error(f"Excel generation failed: {exc}")


@tool(
    description=(
        "Generate a real .docx Word document from markdown-like content. "
        "Use when the user wants a Word file, memo, report, or formatted handout."
    )
)
def tool_generate_word_document(
    markdown_content: str,
    title: str = "Wiii Document",
) -> str:
    """Create a Word document for the user."""
    content = (markdown_content or "").strip()
    if not content:
        return _json_error("markdown_content must not be empty")

    try:
        from docx import Document

        output_dir = _shared_generated_dir()
        filename = f"{_slugify(title, 'wiii-document')}_{_timestamp()}.docx"
        filepath = output_dir / filename

        document = Document()
        document.core_properties.title = title
        document.add_heading(title, level=0)
        _write_markdown_to_docx(document, content)
        document.save(str(filepath))
        metadata = _build_artifact_metadata(filepath)
        metadata["preview"] = content[:240]
        _emit_generated_artifact(
            artifact_type="document",
            title=title,
            filename=filename,
            content=content,
            language="markdown",
            metadata=metadata,
        )

        return json.dumps(
            {
                "file_path": str(filepath),
                "file_url": metadata["file_url"],
                "filename": filename,
                "format": "docx",
                "title": title,
                "preview": content[:240],
            },
            ensure_ascii=False,
        )
    except Exception as exc:
        logger.error("[OUTPUT_TOOLS] Word generation failed: %s", exc)
        return _json_error(f"Word generation failed: {exc}")


def get_output_generation_tools() -> list:
    """Return the deterministic file-generation tools."""
    return [
        tool_generate_html_file,
        tool_generate_excel_file,
        tool_generate_word_document,
    ]


def init_output_generation_tools() -> None:
    """Register deterministic output tools with the global registry."""
    registry = get_tool_registry()
    registry.register(
        tool_generate_html_file,
        ToolCategory.UTILITY,
        ToolAccess.WRITE,
        description="Generate exportable HTML pages and landing pages",
    )
    registry.register(
        tool_generate_excel_file,
        ToolCategory.UTILITY,
        ToolAccess.WRITE,
        description="Generate generic Excel spreadsheets from JSON rows",
    )
    registry.register(
        tool_generate_word_document,
        ToolCategory.UTILITY,
        ToolAccess.WRITE,
        description="Generate Word documents from markdown-like content",
    )
    logger.info("Output generation tools registered")
