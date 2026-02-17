import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx
d = docx.Document(r'C:\Users\Admin\Downloads\Vande\thinking-ux-option-a-report.docx')
out = r'E:\Sach\Sua\AI_v1\docx_output.txt'
with open(out, 'w', encoding='utf-8') as f:
    for p in d.paragraphs:
        f.write(p.text + '\n')
    f.write('\n=== TABLES ===\n')
    for i, table in enumerate(d.tables):
        f.write(f'\n--- Table {i+1} ---\n')
        for row in table.rows:
            f.write(' | '.join(cell.text for cell in row.cells) + '\n')
print(f'Wrote to {out}')
